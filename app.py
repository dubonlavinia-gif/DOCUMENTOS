import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

st.set_page_config(page_title="Generador de Informes", page_icon="📑")

st.title("📑 Generador de Informes (Estructura Fija)")
st.info("Este generador incluirá todas las secciones en el Word, aunque las dejes vacías.")

# --- 1. DATOS GENERALES ---
st.subheader("1. Encabezado")
col1, col2 = st.columns(2)
with col1:
    tipo_informe = st.selectbox("Tipo:", ["INFORME SOCIAL", "INFORME DE VISITA", "CONSTANCIA", "OFICIO"])
    fecha = st.date_input("Fecha:", datetime.date.today())
    ciudad = st.text_input("Ciudad:", "Tegucigalpa, M.D.C.")
with col2:
    destinatario = st.text_input("Dirigido a:", "Jefa de Departamento")
    asunto = st.text_input("Asunto:", "Informe de caso")

# --- 2. CUERPO DEL DOCUMENTO ---
st.subheader("2. Contenido")
intro = st.text_area("1. Antecedentes / Introducción:", "Escribe aquí los antecedentes del caso...")
desarrollo = st.text_area("2. Desarrollo / Hallazgos:", "Escribe aquí lo que encontraste o realizaste...")
conclusion = st.text_area("3. Conclusiones / Recomendaciones:", "Escribe aquí tus recomendaciones profesionales...")

# --- 3. FIRMA ---
st.subheader("3. Firma")
nombre_pro = st.text_input("Tu Nombre:", "Lic. Tu Nombre")
cargo_pro = st.text_input("Tu Cargo:", "Trabajadora Social")

# --- FUNCIÓN GENERADORA ---
def crear_documento(tipo, fecha, ciudad, dest, asunto, intro, des, conc, nom, cargo):
    doc = Document()
    
    # 1. FECHA Y CIUDAD
    p = doc.add_paragraph(f"{ciudad}, {fecha.strftime('%d/%m/%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph() # Espacio

    # 2. TÍTULO CENTRADO
    titulo = doc.add_paragraph(tipo)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].bold = True
    titulo.runs[0].font.size = Pt(14)

    doc.add_paragraph() # Espacio

    # 3. MEMBRETE
    p_head = doc.add_paragraph()
    p_head.add_run("PARA: \t").bold = True
    p_head.add_run(f"{dest}\n")
    p_head.add_run("ASUNTO: \t").bold = True
    p_head.add_run(f"{asunto}")
    
    doc.add_paragraph("__________________________________________________________________________")

    # 4. SECCIONES (Aparecerán SIEMPRE)
    # Antecedentes
    doc.add_heading('1. ANTECEDENTES', level=2)
    doc.add_paragraph(intro)
    
    # Desarrollo
    doc.add_heading('2. DESARROLLO', level=2)
    doc.add_paragraph(des)
    
    # Conclusiones
    doc.add_heading('3. CONCLUSIONES', level=2)
    doc.add_paragraph(conc)

    doc.add_paragraph("\n\n\n")

    # 5. FIRMA
    firma = doc.add_paragraph(f"__________________________\n{nom}\n{cargo}")
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- BOTÓN DE DESCARGA ---
# Pasamos las variables explícitamente para asegurar que se guarden
archivo_listo = crear_documento(
    tipo_informe, fecha, ciudad, destinatario, asunto, 
    intro, desarrollo, conclusion, nombre_pro, cargo_pro
)

st.download_button(
    label="⬇️ Descargar Informe Completo",
    data=archivo_listo,
    file_name="Nuevo_Informe.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)